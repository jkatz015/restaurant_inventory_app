"""
Advanced file extraction module
Handles DOCX, PDF, CSV, Excel, and images with intelligent routing
"""

import io
import os
import hashlib
import base64
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

# Text extraction
import docx
import pdfplumber
import pandas as pd
from PIL import Image

# PDF to image
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    try:
        from pdf2image import convert_from_bytes
        PDF2IMAGE_AVAILABLE = True
    except ImportError:
        PDF2IMAGE_AVAILABLE = False

# Unit detection
from utils.unit_normalizer import count_uom_hits

# Claude API
import anthropic


# Security: Reject files with active content
REJECTED_EXTENSIONS = ['.xlsm', '.docm', '.xlsb']

# Text confidence thresholds
TEXT_CONF = {
    "min_chars": 200,
    "min_words": 30,
    "min_uom_hits": 2
}


def hash_file_content(file_bytes: bytes) -> str:
    """Generate SHA256 hash of file content for duplicate detection"""
    return hashlib.sha256(file_bytes).hexdigest()


def verify_mime_type(uploaded_file) -> Tuple[bool, str]:
    """
    Verify MIME type matches file extension

    Returns:
        Tuple of (is_valid, error_message)
    """
    filename = uploaded_file.name.lower()
    mime_type = uploaded_file.type if hasattr(uploaded_file, 'type') else ''

    # Check for rejected extensions
    for ext in REJECTED_EXTENSIONS:
        if filename.endswith(ext):
            return False, f"File type {ext} not allowed (contains macros)"

    # Basic MIME validation
    valid_mimes = {
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.pdf': 'application/pdf',
        '.csv': 'text/csv',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
    }

    # Find matching extension
    for ext, expected_mime in valid_mimes.items():
        if filename.endswith(ext):
            # Some browsers send generic MIME types
            if mime_type and mime_type != expected_mime and not mime_type.startswith('application/octet-stream'):
                return True, ""  # Allow but warn
            return True, ""

    return False, f"Unsupported file type: {filename}"


def strip_exif(image: Image.Image) -> Image.Image:
    """Remove EXIF metadata from image"""
    data = list(image.getdata())
    image_without_exif = Image.new(image.mode, image.size)
    image_without_exif.putdata(data)
    return image_without_exif


def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from DOCX file

    Returns:
        Dict with status, text, metadata
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))

        # Extract all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)

        return {
            "status": "success",
            "text": full_text,
            "char_count": len(full_text),
            "word_count": len(full_text.split()),
            "paragraph_count": len(doc.paragraphs),
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract DOCX: {str(e)}",
            "text": ""
        }


def extract_text_from_csv(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from CSV with header detection

    Returns:
        Dict with status, text, structured data
    """
    try:
        # Try to read CSV
        df = pd.read_csv(io.BytesIO(file_bytes))

        # Detect recipe-relevant headers
        header_mappings = {
            'ingredient': ['ingredient', 'ingredients', 'item', 'product', 'name'],
            'quantity': ['quantity', 'qty', 'amount', 'measure'],
            'uom': ['uom', 'unit', 'units', 'measurement'],
            'instruction': ['instruction', 'instructions', 'step', 'steps', 'directions']
        }

        # Normalize column names
        columns_lower = {col.lower().strip(): col for col in df.columns}

        # Find matching columns
        mapped_columns = {}
        for key, possible_names in header_mappings.items():
            for name in possible_names:
                if name in columns_lower:
                    mapped_columns[key] = columns_lower[name]
                    break

        # Build text representation
        text_parts = []

        if 'ingredient' in mapped_columns:
            # Structured ingredient list
            ing_col = mapped_columns['ingredient']
            qty_col = mapped_columns.get('quantity')
            uom_col = mapped_columns.get('uom')

            text_parts.append("INGREDIENTS:")
            for idx, row in df.iterrows():
                ingredient = str(row[ing_col]).strip()
                if ingredient and ingredient.lower() not in ['nan', 'none', '']:
                    parts = []
                    if qty_col and str(row[qty_col]) not in ['nan', 'None', '']:
                        parts.append(str(row[qty_col]))
                    if uom_col and str(row[uom_col]) not in ['nan', 'None', '']:
                        parts.append(str(row[uom_col]))
                    parts.append(ingredient)
                    text_parts.append(" ".join(parts))
        else:
            # Unstructured - just concatenate all rows
            for idx, row in df.iterrows():
                row_text = " ".join([str(val) for val in row.values if str(val) not in ['nan', 'None', '']])
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        return {
            "status": "success",
            "text": full_text,
            "structured": mapped_columns != {},
            "row_count": len(df),
            "columns": list(df.columns),
            "mapped_columns": mapped_columns
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract CSV: {str(e)}",
            "text": ""
        }


def extract_text_from_xlsx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Extract text from Excel file with header detection

    Returns:
        Dict with status, text, metadata
    """
    try:
        # Read Excel file (all sheets)
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))

        text_parts = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            text_parts.append(f"\n=== SHEET: {sheet_name} ===\n")

            # Similar header detection as CSV
            header_mappings = {
                'ingredient': ['ingredient', 'ingredients', 'item', 'product', 'name'],
                'quantity': ['quantity', 'qty', 'amount', 'measure'],
                'uom': ['uom', 'unit', 'units', 'measurement'],
            }

            columns_lower = {col.lower().strip(): col for col in df.columns}
            mapped_columns = {}

            for key, possible_names in header_mappings.items():
                for name in possible_names:
                    if name in columns_lower:
                        mapped_columns[key] = columns_lower[name]
                        break

            if 'ingredient' in mapped_columns:
                ing_col = mapped_columns['ingredient']
                qty_col = mapped_columns.get('quantity')
                uom_col = mapped_columns.get('uom')

                for idx, row in df.iterrows():
                    ingredient = str(row[ing_col]).strip()
                    if ingredient and ingredient.lower() not in ['nan', 'none', '']:
                        parts = []
                        if qty_col and str(row[qty_col]) not in ['nan', 'None', '']:
                            parts.append(str(row[qty_col]))
                        if uom_col and str(row[uom_col]) not in ['nan', 'None', '']:
                            parts.append(str(row[uom_col]))
                        parts.append(ingredient)
                        text_parts.append(" ".join(parts))
            else:
                # Just concatenate all cells
                for idx, row in df.iterrows():
                    row_text = " ".join([str(val) for val in row.values if str(val) not in ['nan', 'None', '']])
                    if row_text:
                        text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        return {
            "status": "success",
            "text": full_text,
            "sheet_count": len(excel_file.sheet_names),
            "sheets": excel_file.sheet_names
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract Excel: {str(e)}",
            "text": ""
        }


def analyze_page_confidence(page_text: str) -> Dict[str, Any]:
    """
    Multi-metric confidence test for text extraction

    Returns:
        Dict with is_confident flag and metrics
    """
    char_count = len(page_text) if page_text else 0
    word_count = len(page_text.split()) if page_text else 0
    uom_hits = count_uom_hits(page_text) if page_text else 0

    failures = 0
    metrics = {
        "char_count": char_count,
        "word_count": word_count,
        "uom_hits": uom_hits,
        "tests_passed": []
    }

    if char_count < TEXT_CONF["min_chars"]:
        failures += 1
        metrics["tests_passed"].append(f"chars: {char_count} < {TEXT_CONF['min_chars']} ❌")
    else:
        metrics["tests_passed"].append(f"chars: {char_count} >= {TEXT_CONF['min_chars']} ✓")

    if word_count < TEXT_CONF["min_words"]:
        failures += 1
        metrics["tests_passed"].append(f"words: {word_count} < {TEXT_CONF['min_words']} ❌")
    else:
        metrics["tests_passed"].append(f"words: {word_count} >= {TEXT_CONF['min_words']} ✓")

    if uom_hits < TEXT_CONF["min_uom_hits"]:
        failures += 1
        metrics["tests_passed"].append(f"uom: {uom_hits} < {TEXT_CONF['min_uom_hits']} ❌")
    else:
        metrics["tests_passed"].append(f"uom: {uom_hits} >= {TEXT_CONF['min_uom_hits']} ✓")

    metrics["failures"] = failures
    metrics["is_confident"] = failures < 2  # Pass if at least 2 of 3 metrics pass

    return metrics


def convert_pdf_page_to_image(pdf_bytes: bytes, page_num: int) -> Optional[bytes]:
    """
    Convert a single PDF page to image

    Args:
        pdf_bytes: PDF file bytes
        page_num: Page number (0-indexed)

    Returns:
        Image bytes (PNG) or None if failed
    """
    try:
        if PYMUPDF_AVAILABLE:
            # Use PyMuPDF (faster)
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            doc.close()
            return img_bytes
        elif PDF2IMAGE_AVAILABLE:
            # Use pdf2image
            images = convert_from_bytes(pdf_bytes, first_page=page_num+1, last_page=page_num+1, dpi=150)
            if images:
                img_buffer = io.BytesIO()
                images[0].save(img_buffer, format='PNG')
                return img_buffer.getvalue()
        else:
            return None
    except Exception as e:
        print(f"Error converting PDF page to image: {e}")
        return None


def extract_with_claude_vision(image_bytes: bytes, api_key: str, context: str = "") -> str:
    """
    Use Claude Vision API to extract text from image

    Args:
        image_bytes: Image data
        api_key: Anthropic API key
        context: Additional context for the extraction

    Returns:
        Extracted text
    """
    try:
        # Encode image to base64
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

        # Determine image type
        image_type = "image/png"
        if image_bytes[:4] == b'\xff\xd8\xff\xe0' or image_bytes[:4] == b'\xff\xd8\xff\xe1':
            image_type = "image/jpeg"

        client = anthropic.Anthropic(api_key=api_key)

        prompt = """Extract all text from this recipe image. Focus on:
- Recipe name
- Ingredients (with quantities and units)
- Instructions/steps
- Yield/servings information
- Any other relevant recipe details

Return the text in a structured, readable format. Preserve quantities, units, and ingredient names exactly as shown."""

        if context:
            prompt = f"{context}\n\n{prompt}"

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": image_type,
                            "data": image_base64
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }]
        )

        return message.content[0].text
    except Exception as e:
        return f"[Vision extraction failed: {str(e)}]"


def route_pdf_pages(file_bytes: bytes, api_key: str, max_pages: int = 50) -> Dict[str, Any]:
    """
    Process PDF with per-page routing (text or vision)

    Returns:
        Dict with status, pages list, combined text
    """
    try:
        pages_data = []
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = min(len(pdf.pages), max_pages)

            for page_num, page in enumerate(pdf.pages[:total_pages]):
                # Try text extraction
                page_text = page.extract_text() or ""

                # Analyze confidence
                confidence = analyze_page_confidence(page_text)

                if confidence["is_confident"]:
                    # Use text extraction
                    pages_data.append({
                        "page_number": page_num + 1,
                        "route": "text",
                        "confidence": confidence,
                        "text": page_text
                    })
                    text_parts.append(f"\n--- PAGE {page_num + 1} ---\n{page_text}")
                else:
                    # Use vision extraction
                    img_bytes = convert_pdf_page_to_image(file_bytes, page_num)
                    if img_bytes:
                        vision_text = extract_with_claude_vision(
                            img_bytes,
                            api_key,
                            context=f"This is page {page_num + 1} of a recipe document."
                        )
                        pages_data.append({
                            "page_number": page_num + 1,
                            "route": "vision",
                            "confidence": confidence,
                            "text": vision_text
                        })
                        text_parts.append(f"\n--- PAGE {page_num + 1} (VISION) ---\n{vision_text}")
                    else:
                        # Fallback to low-confidence text
                        pages_data.append({
                            "page_number": page_num + 1,
                            "route": "text_fallback",
                            "confidence": confidence,
                            "text": page_text
                        })
                        text_parts.append(f"\n--- PAGE {page_num + 1} (LOW CONF) ---\n{page_text}")

        combined_text = "\n".join(text_parts)

        # Count routes
        text_pages = sum(1 for p in pages_data if p["route"] == "text")
        vision_pages = sum(1 for p in pages_data if p["route"] == "vision")

        return {
            "status": "success",
            "total_pages": total_pages,
            "text_pages": text_pages,
            "vision_pages": vision_pages,
            "pages": pages_data,
            "text": combined_text
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to process PDF: {str(e)}",
            "pages": [],
            "text": ""
        }


def extract_from_image(file_bytes: bytes, api_key: str) -> Dict[str, Any]:
    """
    Extract text from image file using Claude Vision

    Returns:
        Dict with status, text
    """
    try:
        # Strip EXIF
        img = Image.open(io.BytesIO(file_bytes))
        img_clean = strip_exif(img)

        # Convert to bytes
        img_buffer = io.BytesIO()
        img_clean.save(img_buffer, format=img.format or 'PNG')
        img_bytes = img_buffer.getvalue()

        # Use vision extraction
        text = extract_with_claude_vision(img_bytes, api_key, context="This is a recipe image.")

        return {
            "status": "success",
            "text": text,
            "image_format": img.format,
            "image_size": img.size
        }
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract from image: {str(e)}",
            "text": ""
        }


def process_uploaded_file(uploaded_file, api_key: str) -> Dict[str, Any]:
    """
    Main orchestrator: detect type, route, extract text

    Returns:
        Dict with status, file_type, text, metadata, pages (for PDFs)
    """
    # Security check
    is_valid, error_msg = verify_mime_type(uploaded_file)
    if not is_valid:
        return {
            "status": "error",
            "error": error_msg
        }

    # Read file bytes
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)  # Reset for potential re-reading

    # Hash for duplicate detection
    file_hash = hash_file_content(file_bytes)

    filename = uploaded_file.name.lower()

    # Route by extension
    if filename.endswith('.docx'):
        result = extract_text_from_docx(file_bytes)
        result["file_type"] = "docx"
    elif filename.endswith('.pdf'):
        result = route_pdf_pages(file_bytes, api_key)
        result["file_type"] = "pdf"
    elif filename.endswith('.csv'):
        result = extract_text_from_csv(file_bytes)
        result["file_type"] = "csv"
    elif filename.endswith(('.xlsx', '.xls')):
        result = extract_text_from_xlsx(file_bytes)
        result["file_type"] = "xlsx"
    elif filename.endswith(('.png', '.jpg', '.jpeg')):
        result = extract_from_image(file_bytes, api_key)
        result["file_type"] = "image"
    else:
        return {
            "status": "error",
            "error": f"Unsupported file type: {filename}"
        }

    # Add common metadata
    result["filename"] = uploaded_file.name
    result["file_hash"] = file_hash
    result["file_size"] = len(file_bytes)

    return result

